"""
H2READY Toolkit - generazione dei documenti (PDF e Word).

Entrambi i formati partono dallo stesso testo in Markdown prodotto da contenuti.py,
così i due documenti restano allineati.
"""

import io
import os
import re
from datetime import date

from fpdf import FPDF
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

LOGO = "logo_h2ready.png"
FONT_DIR = "fonts"
BLU = (0, 51, 153)
BLU_HEX = "003399"
GRIGIO = (110, 110, 110)

# =============================================================================
# PARSING DEL MARKDOWN
# =============================================================================

def blocchi_markdown(md: str):
    """Trasforma il markdown in una sequenza di blocchi tipizzati."""
    righe = md.split("\n")
    i = 0
    while i < len(righe):
        stripped = righe[i].strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            tabella = []
            while i < len(righe) and righe[i].strip().startswith("|"):
                celle = [c.strip() for c in righe[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in celle if c):
                    tabella.append(celle)
                i += 1
            if tabella:
                yield ("tabella", tabella)
            continue

        if stripped == "<<<PAGINA>>>":
            yield ("pagina", None)
        elif re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped):
            yield ("riga", None)
        elif stripped.startswith("####"):
            yield ("h4", stripped.lstrip("#").strip())
        elif stripped.startswith("###"):
            yield ("h3", stripped.lstrip("#").strip())
        elif stripped.startswith("##"):
            yield ("h2", stripped.lstrip("#").strip())
        elif stripped.startswith("#"):
            yield ("h1", stripped.lstrip("#").strip())
        elif stripped.startswith(">"):
            nota = []
            while i < len(righe) and righe[i].strip().startswith(">"):
                nota.append(righe[i].strip().lstrip(">").strip())
                i += 1
            yield ("nota", " ".join(p for p in nota if p))
            continue
        elif re.match(r"^([-*+]|\d{1,2}[.)])\s+", stripped):
            marc = re.match(r"^([-*+]|\d{1,2}[.)])\s+", stripped).group(1)
            testo = re.sub(r"^([-*+]|\d{1,2}[.)])\s+", "", stripped)
            yield ("elenco", (marc if marc not in "-*+" else "-", testo))
        else:
            paragrafo = [stripped]
            i += 1
            while i < len(righe):
                seguente = righe[i].strip()
                if (not seguente or seguente.startswith(("|", "#", ">"))
                        or re.match(r"^([-*+]|\d{1,2}[.)])\s+", seguente)
                        or re.fullmatch(r"-{3,}|_{3,}|\*{3,}", seguente)):
                    break
                paragrafo.append(seguente)
                i += 1
            yield ("paragrafo", " ".join(paragrafo))
            continue
        i += 1


def spezza_grassetto(testo: str):
    """Scompone il testo in (frammento, grassetto, url).

    I link markdown vengono prima messi da parte con un segnaposto: se si
    dividesse subito la stringa sui link, un collegamento scritto dentro un
    blocco in grassetto lascerebbe asterischi orfani nel documento.
    """
    link = []

    def _cattura(m):
        link.append((m.group(1), m.group(2)))
        return f"\x00{len(link) - 1}\x00"

    testo = re.sub(r"\[([^\]]+)\]\(((?:https?://|mailto:)[^)]+)\)", _cattura, testo)

    parti = []
    for pezzo in re.split(r"(\*\*.+?\*\*)", testo):
        if not pezzo:
            continue
        grassetto = pezzo.startswith("**") and pezzo.endswith("**")
        if grassetto:
            pezzo = pezzo[2:-2]
        for frammento in re.split(r"\x00(\d+)\x00", pezzo):
            if frammento == "":
                continue
            if frammento.isdigit() and int(frammento) < len(link):
                etichetta, url = link[int(frammento)]
                # un **grassetto** scritto dentro l'etichetta di un collegamento
                # non è interpretabile: il link ha già una sua evidenza grafica
                pulita = etichetta.replace("**", "")
                parti.append((pulita, grassetto or etichetta != pulita, url))
            else:
                parti.append((frammento.replace("**", ""), grassetto, None))
    return parti



# =============================================================================
# INDICE
# =============================================================================

def _registra(pdf, titolo, livello):
    """Annota un titolo per l'indice, con la pagina in cui compare."""
    if not hasattr(pdf, "voci_indice"):
        pdf.voci_indice = []
    # si registra la pagina che ospiterà davvero il titolo: se lo spazio residuo
    # non basta, il titolo scivola alla pagina successiva e l'indice sbaglierebbe
    if pdf.get_y() > pdf.h - 45:
        pdf.add_page()
    pdf.voci_indice.append((livello, str(titolo), pdf.page_no()))


def _scrivi_indice(pdf, voci, offset):
    """Compone la pagina di indice. offset è il numero di pagine che l'indice
    stesso occupa, da sommare ai numeri raccolti nella prima passata."""
    pdf.font("B", 17)
    pdf.set_text_color(*BLU)
    pdf.multi_cell(0, 11, pdf.txt("Indice"))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    larghezza = pdf.w - pdf.l_margin - pdf.r_margin
    for livello, titolo, pagina in voci:
        # 0 = passo, 1 = titolo di sezione, 2 = sottosezione
        if livello == 0:
            pdf.ln(3)
            pdf.font("B", 11)
            pdf.set_text_color(*BLU)
            rientro = 0
        elif livello == 1:
            pdf.font("B", 10)
            pdf.set_text_color(0, 0, 0)
            rientro = 5
        else:
            pdf.font("", 9.5)
            pdf.set_text_color(70, 78, 90)
            rientro = 11

        numero_pagina = str(pagina + offset)
        pdf.set_x(pdf.l_margin + rientro)
        testo = pdf.txt(titolo)
        larghezza_num = pdf.get_string_width(numero_pagina)
        disponibile = larghezza - rientro - larghezza_num - 4

        # il titolo si accorcia se non entra: l'indice deve restare su una riga
        while pdf.get_string_width(testo) > disponibile and len(testo) > 12:
            testo = testo[:-2]
        if pdf.get_string_width(pdf.txt(titolo)) > disponibile:
            testo += "..."

        pdf.cell(pdf.get_string_width(testo), 6, testo)
        puntini = ""
        larghezza_punto = pdf.get_string_width(".")
        spazio = larghezza - rientro - pdf.get_string_width(testo) - larghezza_num - 2
        if larghezza_punto > 0:
            pdf.set_text_color(170, 176, 186)
            puntini = "." * max(0, int(spazio / larghezza_punto))
            pdf.cell(spazio, 6, puntini)
            pdf.set_text_color(0, 0, 0)
        pdf.cell(larghezza_num + 2, 6, numero_pagina, align="R",
                 new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)


# =============================================================================
# PDF
# =============================================================================

def pulisci(testo) -> str:
    if not isinstance(testo, str):
        testo = str(testo)
    sost = {"\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
            "\u2013": "-", "\u2014": "-", "\u2026": "...", "\u20ac": "Euro",
            "\u2082": "2", "\u2083": "3", "\u2705": "-", "\u2022": "-",
            "\u00b7": "-", "\u2192": "->", "\u2264": "<=", "\u2265": ">="}
    for a, b in sost.items():
        testo = testo.replace(a, b)
    # emoji e simboli fuori dal latin-1: meglio toglierli che stampare "?"
    testo = re.sub(
        "[\U0001F000-\U0001FAFF\u2190-\u21FF\u2300-\u27BF\uFE0F\u2B00-\u2BFF]",
        "", testo)
    testo = re.sub(r" {2,}", " ", testo)
    return testo.encode("latin-1", "replace").decode("latin-1")


class H2ReadyPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.pagina_piena = False
        self.pagine_piene = set()
        self.set_auto_page_break(True, margin=22)
        self.set_margins(20, 20, 20)
        self.unicode = False
        self.italico = True
        self._carica_font()

    def _carica_font(self):
        reg = os.path.join(FONT_DIR, "DejaVuSans.ttf")
        bol = os.path.join(FONT_DIR, "DejaVuSans-Bold.ttf")
        ita = os.path.join(FONT_DIR, "DejaVuSans-Oblique.ttf")
        if os.path.exists(reg) and os.path.exists(bol):
            self.add_font("DejaVu", "", reg)
            self.add_font("DejaVu", "B", bol)
            self.italico = os.path.exists(ita)
            if self.italico:
                self.add_font("DejaVu", "I", ita)
            self.unicode = True

    @property
    def famiglia(self):
        return "DejaVu" if self.unicode else "Arial"

    def txt(self, testo):
        return str(testo) if self.unicode else pulisci(testo)

    def font(self, stile="", size=11):
        if stile == "I" and not self.italico:
            stile = ""
        self.set_font(self.famiglia, stile, size)

    def header(self):
        if self.page_no() == 1 or self.pagina_piena:
            return
        self.pagine_piene.discard(self.page_no())
        if os.path.exists(LOGO):
            self.image(LOGO, 20, 8, 22)
        self.font("B", 8)
        self.set_text_color(*GRIGIO)
        self.set_xy(45, 10)
        self.cell(0, 5, self.txt("H2READY - Interreg VI-A Italia-Slovenia"),
                  new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(*BLU)
        self.line(20, 18, 190, 18)
        self.set_y(28)
        self.set_text_color(0, 0, 0)

    def footer(self):
        if self.page_no() in self.pagine_piene:
            return
        self.set_y(-15)
        self.font("I", 8)
        self.set_text_color(*GRIGIO)
        self.cell(0, 10, self.txt(f"Pagina {self.page_no()}  |  Action Plan H2READY"),
                  align="C")


def _unita_di_parola(parti):
    """Raggruppa i frammenti in parole intere.

    Serve perché un cambio di stile a metà parola — come in "l'**elettrificazione**" —
    farebbe perdere a fpdf la cognizione che i due pezzi sono attaccati, e la parola
    verrebbe spezzata a fine riga in un punto qualsiasi.
    Ogni unità è una lista di (testo, grassetto, url) da scrivere senza interruzione.
    """
    unita, corrente, attaccato = [], [], False
    for testo, grassetto, url in parti:
        pezzi = re.split(r"(\s+)", testo)
        for pezzo in pezzi:
            if pezzo == "":
                continue
            if pezzo.isspace():
                if corrente:
                    unita.append(corrente)
                    corrente = []
                unita.append(None)          # separatore: spazio
                attaccato = False
                continue
            corrente.append((pezzo, grassetto, url))
            attaccato = True
    if corrente:
        unita.append(corrente)
    return unita


def _pdf_inline(pdf, testo, size, h):
    """Scrive un paragrafo andando a capo solo fra una parola e l'altra."""
    limite = pdf.w - pdf.r_margin - 2.0

    def larghezza(gruppo):
        totale = 0.0
        for frammento, grassetto, _ in gruppo:
            pdf.font("B" if grassetto else "", size)
            totale += pdf.get_string_width(pdf.txt(frammento))
        return totale

    pdf.font("", size)
    spazio = pdf.get_string_width(" ")
    unita = [u for u in _unita_di_parola(spezza_grassetto(testo))]
    parole = [u for u in unita if u is not None]
    primo = True

    for i, gruppo in enumerate(parole):
        larga = larghezza(gruppo)
        prefisso = 0.0 if primo else spazio
        if not primo:
            # lo spazio si scrive solo se anche la parola che segue entra nella
            # riga: altrimenti resterebbe appeso in fondo e la parola verrebbe
            # spezzata dal ritorno a capo automatico di fpdf
            if pdf.get_x() + prefisso + larga > limite:
                pdf.ln(h)
                pdf.set_x(pdf.l_margin)
            else:
                pdf.font("", size)
                pdf.write(h, " ")
        for frammento, grassetto, url in gruppo:
            pdf.font("B" if grassetto else "", size)
            if url:
                pdf.set_text_color(*BLU)
                pdf.write(h, pdf.txt(frammento), link=url)
                pdf.set_text_color(0, 0, 0)
            else:
                pdf.write(h, pdf.txt(frammento))
        primo = False
    pdf.ln(h)


def _pdf_tabella(pdf, dati):
    """Tabella a larghezza adattiva: usa l'API table di fpdf2, che gestisce
    il ritorno a capo nelle celle e l'interruzione di pagina."""
    n = max(len(r) for r in dati)
    dati = [list(r) + [""] * (n - len(r)) for r in dati]

    lunghezze = []
    for j in range(n):
        massimo = max(len(str(r[j])) for r in dati)
        # esponente < 1: le colonne lunghe crescono, ma non schiacciano le altre
        lunghezze.append(max(massimo, 6) ** 0.65)
    totale = sum(lunghezze)
    pesi = [l / totale for l in lunghezze]

    # nessuna colonna può essere così stretta da spezzare una parola dell'intestazione
    utile = pdf.w - pdf.l_margin - pdf.r_margin
    for j in range(n):
        parola = max(str(dati[0][j]).split() or [""], key=len)
        minimo = (len(parola) * 1.85 + 4) / utile
        pesi[j] = max(pesi[j], min(minimo, 0.30))
    somma = sum(pesi)
    pesi = [p / somma for p in pesi]

    pdf.font("", 9)
    pdf.set_draw_color(190, 198, 214)
    # lo stato di riempimento resta quello impostato altrove (le pagine divisorie
    # lo lasciano sul blu pieno): va riportato al bianco, altrimenti le celle
    # escono con il fondo scuro e il testo diventa illeggibile
    pdf.set_fill_color(255, 255, 255)
    pdf.set_text_color(0, 0, 0)
    with pdf.table(col_widths=tuple(pesi),
                   text_align=tuple(_allinea(dati, j, n) for j in range(n)),
                   line_height=5,
                   padding=1.6,
                   headings_style=_intestazione_tabella(pdf),
                   cell_fill_mode="NONE") as tabella:
        for i, riga in enumerate(dati):
            fila = tabella.row()
            for cella in riga:
                fila.cell(pdf.txt(_senza_grassetto(cella)))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)


def _allinea(dati, j, n) -> str:
    """Numeri e valori brevi a destra, testo discorsivo a sinistra."""
    if j == 0:
        return "LEFT"
    valori = [str(r[j]) for r in dati[1:] if j < len(r)]
    if valori and max(len(v) for v in valori) > 28:
        return "LEFT"
    return "RIGHT"


def _senza_grassetto(testo) -> str:
    """Nelle celle il markdown non viene interpretato: resta il solo testo."""
    pulito = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", str(testo))
    return re.sub(r"\*\*(.+?)\*\*", r"\1", pulito)


def _intestazione_tabella(pdf):
    """Intestazione su fondo chiaro: il testo bianco su blu pieno risulta poco
    leggibile in stampa e nelle esportazioni a bassa risoluzione."""
    from fpdf.fonts import FontFace
    return FontFace(emphasis="BOLD", color=BLU, fill_color=(226, 232, 243))


def scrivi_markdown_pdf(pdf, md: str):
    for tipo, contenuto in blocchi_markdown(md):
        if tipo == "tabella":
            _pdf_tabella(pdf, contenuto)
        elif tipo == "pagina":
            if pdf.get_y() > pdf.t_margin + 30:
                pdf.add_page()
        elif tipo == "riga":
            pdf.ln(2)
            pdf.set_draw_color(200, 205, 215)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(4)
        elif tipo == "h1":
            pdf.ln(2); pdf.font("B", 17); pdf.set_text_color(*BLU)
            _registra(pdf, contenuto, 1)
            pdf.multi_cell(0, 11, pdf.txt(contenuto))
            pdf.set_text_color(0, 0, 0); pdf.ln(3)
        elif tipo == "h2":
            pdf.ln(4); pdf.font("B", 14); pdf.set_text_color(*BLU)
            _registra(pdf, contenuto, 1)
            pdf.multi_cell(0, 9, pdf.txt(contenuto))
            pdf.set_draw_color(*BLU)
            pdf.line(pdf.l_margin, pdf.get_y() + 1, pdf.l_margin + 35, pdf.get_y() + 1)
            pdf.set_text_color(0, 0, 0); pdf.ln(4)
        elif tipo == "h3":
            pdf.ln(3); pdf.font("B", 12); pdf.set_text_color(*BLU)
            _registra(pdf, contenuto, 2)
            pdf.multi_cell(0, 8, pdf.txt(contenuto))
            pdf.set_text_color(0, 0, 0); pdf.ln(1)
        elif tipo == "h4":
            pdf.ln(2); pdf.font("B", 11); pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 7, pdf.txt(contenuto))
            pdf.set_text_color(0, 0, 0); pdf.ln(1)
        elif tipo == "nota":
            y0 = pdf.get_y()
            pagina0 = pdf.page_no()
            pdf.set_fill_color(240, 243, 250); pdf.font("I", 10)
            pdf.set_x(pdf.l_margin + 4)
            pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 4, 6,
                           pdf.txt(contenuto), fill=True)
            # se la nota è andata a capo pagina, y0 appartiene alla pagina
            # precedente: disegnarne la barra da lì traccerebbe una riga verticale
            # lunga quanto il foglio, accanto a testo che non c'entra nulla
            if pdf.page_no() == pagina0:
                pdf.set_draw_color(*BLU); pdf.set_line_width(1)
                pdf.line(pdf.l_margin + 1, y0, pdf.l_margin + 1, pdf.get_y())
                pdf.set_line_width(0.2)
            pdf.ln(3)
        elif tipo == "elenco":
            marcatore, testo = contenuto
            pdf.font("B", 11)
            pdf.cell(7, 6, pdf.txt(marcatore))
            pdf.set_x(pdf.l_margin + 7)
            pdf.set_left_margin(pdf.get_x())
            _pdf_inline(pdf, testo, 11, 6)
            pdf.set_left_margin(20); pdf.ln(1)
        else:
            _pdf_inline(pdf, contenuto, 11, 6.5)
            pdf.ln(2)


def _pdf_divisoria(pdf, occhiello, titolo, sottotitolo=""):
    pdf.pagina_piena = True
    pdf.add_page()
    pdf.pagine_piene.add(pdf.page_no())
    pdf.set_fill_color(*BLU)
    pdf.rect(0, 0, 210, 297, "F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_y(110); pdf.font("B", 13)
    pdf.cell(0, 8, pdf.txt(occhiello.upper()), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_draw_color(255, 255, 255)
    pdf.line(85, pdf.get_y(), 125, pdf.get_y())
    pdf.ln(10); pdf.font("B", 22)
    pdf.multi_cell(0, 12, pdf.txt(titolo), align="C")
    if sottotitolo:
        pdf.ln(6); pdf.font("", 12)
        pdf.multi_cell(0, 7, pdf.txt(sottotitolo), align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.pagina_piena = False


def _pdf_copertina(pdf, comune, livello, profilo):
    pdf.pagina_piena = True
    pdf.add_page()
    pdf.pagine_piene.add(pdf.page_no())
    pdf.set_fill_color(*BLU)
    pdf.rect(0, 0, 12, 297, "F")
    pdf.rect(0, 250, 210, 47, "F")
    if os.path.exists(LOGO):
        pdf.image(LOGO, x=75, y=32, w=60)
    pdf.set_y(105); pdf.font("", 16); pdf.set_text_color(*GRIGIO)
    pdf.cell(0, 10, pdf.txt("ACTION PLAN"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.font("B", 38); pdf.set_text_color(*BLU)
    pdf.cell(0, 20, pdf.txt("H2READY"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(18); pdf.font("B", 24); pdf.set_text_color(20, 20, 20)
    pdf.multi_cell(0, 13, pdf.txt(f"COMUNE DI {comune.upper()}"), align="C")
    pdf.ln(6); pdf.font("", 12); pdf.set_text_color(*GRIGIO)
    pdf.cell(0, 7, pdf.txt(f"Livello di maturità: {livello}   |   Profilo strategico: {profilo}"),
             align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 7, pdf.txt(f"Documento generato il {date.today().strftime('%d/%m/%Y')}"),
             align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_y(262); pdf.set_text_color(255, 255, 255); pdf.font("B", 11)
    pdf.cell(0, 6, pdf.txt("Documento strategico di transizione energetica"),
             align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.font("", 10)
    pdf.cell(0, 6, pdf.txt("Progetto cofinanziato dall'Unione Europea - Interreg VI-A Italia-Slovenia"),
             align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.pagina_piena = False


PASSI = [
    ("Passo 1", "Livello di maturità e profilo strategico",
     ["mat_intro", "mat_dettaglio", "profilo_intro", "profilo_calcolato", "profilo_dettaglio"]),
    ("Passo 2", "Risultato dei percorsi identificati", ["passo2"]),
    ("Passo 3", "Il piano d'azione", ["passo3"]),
]


def _componi(c: dict, voci_indice=None, pagine_indice=0):
    """Costruisce il documento. Se voci_indice è fornito, inserisce l'indice
    dopo la copertina usando quelle voci."""
    pdf = H2ReadyPDF()
    pdf.voci_indice = []
    _pdf_copertina(pdf, c["comune"], c["livello"], c["profilo"] or "n.d.")

    if voci_indice is not None:
        pdf.add_page()
        _scrivi_indice(pdf, voci_indice, pagine_indice)

    pdf.add_page(); scrivi_markdown_pdf(pdf, c["intro"])
    pdf.add_page(); scrivi_markdown_pdf(pdf, c["struttura"])

    for occhiello, titolo, chiavi in PASSI:
        sotto = f"Comune di {c['comune']}" if occhiello == "Passo 1" else ""
        _pdf_divisoria(pdf, occhiello, titolo, sotto)
        pdf.voci_indice.append((0, f"{occhiello} — {titolo}", pdf.page_no()))
        pdf.add_page()
        for k, chiave in enumerate(chiavi):
            testo = c.get(chiave, "")
            if not testo:
                continue
            if k:
                pdf.ln(4)
            scrivi_markdown_pdf(pdf, testo)
    return pdf


def genera_pdf(c: dict) -> bytes:
    """Il documento si compone due volte.

    La prima serve solo a sapere in quale pagina finisce ciascun titolo; la
    seconda inserisce l'indice e ricalcola i numeri tenendo conto delle pagine
    che l'indice stesso occupa. Non esiste modo di saperlo in una passata sola,
    perché la lunghezza dell'indice dipende dal numero di titoli.
    """
    prima = _componi(c)
    voci = prima.voci_indice
    # una pagina di indice ogni 40 voci circa, con un minimo di una
    pagine_indice = max(1, -(-len(voci) // 40))
    return bytes(_componi(c, voci, pagine_indice).output())


# =============================================================================
# WORD
# =============================================================================

def _ombreggia(cella, esadecimale: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), esadecimale)
    cella._tc.get_or_add_tcPr().append(shd)


def _campo(paragrafo, istruzione: str):
    """Inserisce un campo Word (es. numero di pagina)."""
    run = paragrafo.add_run()
    inizio = OxmlElement("w:fldChar"); inizio.set(qn("w:fldCharType"), "begin")
    testo = OxmlElement("w:instrText"); testo.set(qn("xml:space"), "preserve")
    testo.text = istruzione
    fine = OxmlElement("w:fldChar"); fine.set(qn("w:fldCharType"), "end")
    run._r.append(inizio); run._r.append(testo); run._r.append(fine)


def _stili(doc):
    normale = doc.styles["Normal"]
    normale.font.name = "Calibri"
    normale.font.size = Pt(11)
    normale.paragraph_format.space_after = Pt(6)
    for nome, dim in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12),
                      ("Heading 4", 11)):
        stile = doc.styles[nome]
        stile.font.name = "Calibri"
        stile.font.size = Pt(dim)
        stile.font.bold = True
        stile.font.color.rgb = RGBColor(*BLU)


def _aggiungi_link(paragrafo, testo, url, grassetto=False):
    """Un collegamento ipertestuale in Word: python-docx non ha un'API dedicata,
    va costruito a mano nel documento XML."""
    parte = paragrafo.part
    r_id = parte.relate_to(url,
                           "http://schemas.openxmlformats.org/officeDocument/2006/"
                           "relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    prop = OxmlElement("w:rPr")
    colore = OxmlElement("w:color"); colore.set(qn("w:val"), BLU_HEX)
    sottolineato = OxmlElement("w:u"); sottolineato.set(qn("w:val"), "single")
    prop.append(colore); prop.append(sottolineato)
    if grassetto:
        prop.append(OxmlElement("w:b"))
    run.append(prop)
    testo_el = OxmlElement("w:t")
    testo_el.set(qn("xml:space"), "preserve")
    testo_el.text = testo
    run.append(testo_el)
    link.append(run)
    paragrafo._p.append(link)


def _paragrafo_con_grassetto(doc, testo, stile=None):
    p = doc.add_paragraph(style=stile)
    for parte, grassetto, url in spezza_grassetto(testo):
        if url:
            _aggiungi_link(p, parte, url, grassetto)
        else:
            run = p.add_run(parte)
            run.bold = grassetto
    return p


def scrivi_markdown_docx(doc, md: str):
    for tipo, contenuto in blocchi_markdown(md):
        if tipo == "tabella":
            n = max(len(r) for r in contenuto)
            tabella = doc.add_table(rows=0, cols=n)
            tabella.style = "Table Grid"
            tabella.alignment = WD_TABLE_ALIGNMENT.CENTER
            utile = 16.0
            lung = []
            for j in range(n):
                massimo = max(len(str(r[j])) if j < len(r) else 0 for r in contenuto)
                lung.append(max(massimo, 6))
            somma = sum(lung)
            larghezze = [Cm(max(utile * l / somma, 1.6)) for l in lung]
            for i, riga in enumerate(contenuto):
                celle = tabella.add_row().cells
                for j in range(n):
                    testo = _senza_grassetto(riga[j] if j < len(riga) else "")
                    cella = celle[j]
                    cella.width = larghezze[j]
                    p = cella.paragraphs[0]
                    if j and n <= 3:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(testo)
                    run.font.size = Pt(9.5)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(*BLU)
                        _ombreggia(cella, "E2E8F3")

            doc.add_paragraph()
        elif tipo == "pagina":
            doc.add_page_break()
        elif tipo == "riga":
            doc.add_paragraph("_" * 60)
        elif tipo == "h1":
            doc.add_heading(contenuto, level=1)
        elif tipo == "h2":
            doc.add_heading(contenuto, level=2)
        elif tipo == "h3":
            doc.add_heading(contenuto, level=3)
        elif tipo == "h4":
            doc.add_heading(contenuto, level=4)
        elif tipo == "nota":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run(contenuto)
            run.italic = True
            run.font.size = Pt(10)
        elif tipo == "elenco":
            _, testo = contenuto
            _paragrafo_con_grassetto(doc, testo, stile="List Bullet")
        else:
            _paragrafo_con_grassetto(doc, contenuto)


def _docx_copertina(doc, comune, livello, profilo):
    for _ in range(4):
        doc.add_paragraph()
    if os.path.exists(LOGO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(LOGO, width=Cm(6))

    def centrato(testo, dim, grassetto=True, colore=(0, 0, 0)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(testo)
        run.bold = grassetto
        run.font.size = Pt(dim)
        run.font.color.rgb = RGBColor(*colore)
        return p

    doc.add_paragraph()
    centrato("ACTION PLAN", 16, False, GRIGIO)
    centrato("H2READY", 34, True, BLU)
    doc.add_paragraph()
    centrato(f"COMUNE DI {comune.upper()}", 22)
    doc.add_paragraph()
    centrato(f"Livello di maturità: {livello}   |   Profilo strategico: {profilo}",
             11, False, GRIGIO)
    centrato(f"Documento generato il {date.today().strftime('%d/%m/%Y')}", 11, False, GRIGIO)
    for _ in range(4):
        doc.add_paragraph()
    centrato("Documento strategico di transizione energetica", 11, True, BLU)
    centrato("Progetto cofinanziato dall'Unione Europea - Interreg VI-A Italia-Slovenia",
             10, False, GRIGIO)
    doc.add_page_break()


def _docx_divisoria(doc, occhiello, titolo, sottotitolo=""):
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    tabella = doc.add_table(rows=1, cols=1)
    cella = tabella.rows[0].cells[0]
    cella.width = Cm(16)
    _ombreggia(cella, BLU_HEX)

    p = cella.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(occhiello.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)

    p2 = cella.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(titolo)
    run2.bold = True
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(255, 255, 255)

    if sottotitolo:
        p3 = cella.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(sottotitolo)
        run3.font.size = Pt(11)
        run3.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_page_break()


def genera_docx(c: dict) -> bytes:
    doc = Document()
    _stili(doc)

    sezione = doc.sections[0]
    sezione.top_margin = Cm(2.2)
    sezione.bottom_margin = Cm(2.2)
    sezione.left_margin = Cm(2.5)
    sezione.right_margin = Cm(2.5)

    intestazione = sezione.header.paragraphs[0]
    intestazione.text = "H2READY - Interreg VI-A Italia-Slovenia"
    intestazione.runs[0].font.size = Pt(8)
    intestazione.runs[0].font.color.rgb = RGBColor(*GRIGIO)

    pie = sezione.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pie.add_run("Pagina ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*GRIGIO)
    _campo(pie, "PAGE")
    coda = pie.add_run("  |  Action Plan H2READY")
    coda.font.size = Pt(8)
    coda.font.color.rgb = RGBColor(*GRIGIO)

    _docx_copertina(doc, c["comune"], c["livello"], c["profilo"] or "n.d.")
    scrivi_markdown_docx(doc, c["intro"])
    doc.add_page_break()
    scrivi_markdown_docx(doc, c["struttura"])

    for occhiello, titolo, chiavi in PASSI:
        sotto = f"Comune di {c['comune']}" if occhiello == "Passo 1" else ""
        _docx_divisoria(doc, occhiello, titolo, sotto)
        for chiave in chiavi:
            testo = c.get(chiave, "")
            if testo:
                scrivi_markdown_docx(doc, testo)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
